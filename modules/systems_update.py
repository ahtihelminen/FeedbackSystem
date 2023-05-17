from tools import Tools
import json
from docx import Document
from docx.shared import Inches
import os


class SystemsUpdate(Tools):
    def __init__(self, excelSystemsFeedackFile, feedbackDatabase, mode, ship):
        super().__init__()
        self.excelSystemsFeedackFile = excelSystemsFeedackFile
        self.feedbackDatabase = feedbackDatabase
        self.csvSystemsFeedbackFile = None
        self.listSystemsFeedbackFile = None
        self.ship = ship
        self.questionToExclude = ['ID', 'Start time', 'Completion time', 'Email', 'Name', 'valitse littera', 'choose system code', '1000 Ship general design', '3000 Hull', '4000 Interior', '5000 HVAC', '6000 Propulsion', '7000 Machinery', '8000 Deck', '9000 Electric']
        self.mode = mode


    def extractSystems(self, Q_A_dict):
        try:
            systemCode = Q_A_dict['valitse littera']
        except:
            try:
                systemCode = Q_A_dict['choose system code']
            except Exception as e:
                print('Error in extractSystems()', e)
                return False
        finally:
            for question, answer in Q_A_dict.items():
                try:
                    if question.split(' ')[0] == systemCode:
                        return systemCode, answer.split(';')
                except:
                    pass


    def applicableString(self, stringToChange):
        stringToChange = stringToChange.strip()
        stringToChange = stringToChange.replace('\t', ' ')
        stringToChange = stringToChange.replace('\n', ' ')
        return stringToChange


    def updateSystemsDatabase(self):

        self.initializeDatabase()

        self.feedbackDatabase = self.relativeFilepathToAbsolute(self.feedbackDatabase)
        questions = self.listSystemsFeedbackFile[0]
        database = self.readJSON(self.feedbackDatabase)


        with open(self.feedbackDatabase, 'w') as databaseToWrite:
            try:
                
                databaseToUpdate = database
                
                for row in self.listSystemsFeedbackFile[1:]:
                    
                    questionAnswerDict = self.questionAnswerDict(row, questions)
                    questionAnswerDict = self.removeQuestionsWithNoAnswer(questionAnswerDict)
                    systemCode, specificSystems = self.extractSystems(questionAnswerDict)
                    
                    for specificSystem in specificSystems:
                        if specificSystem == '':
                            continue

                        specificSystem = self.applicableString(specificSystem)


                        
                        if specificSystem not in databaseToUpdate['feedbacks']['systems'][systemCode]:
                            databaseToUpdate['feedbacks']['systems'][systemCode][specificSystem] = {}

                        if self.ship not in databaseToUpdate['feedbacks']['systems'][systemCode][specificSystem]:
                            databaseToUpdate['feedbacks']['systems'][systemCode][specificSystem][self.ship] = {}

                        for question, answer in questionAnswerDict.items():

                            if question not in databaseToUpdate['feedbacks']['systems'][systemCode][specificSystem][self.ship]:
                                databaseToUpdate['feedbacks']['systems'][systemCode][specificSystem][self.ship][question] = []
                            
                            if answer not in databaseToUpdate['feedbacks']['systems'][systemCode][specificSystem][self.ship][question]:
                               databaseToUpdate['feedbacks']['systems'][systemCode][specificSystem][self.ship][question].append(answer)

                json.dump(databaseToUpdate, databaseToWrite, indent=4)
            except Exception as e:
                print('Error in updateDatabase()', e)
                json.dump(database, databaseToWrite, indent=4)
                raise TypeError("Error in feedback update:", e)
        databaseToWrite.close()


    def createFeedbackFiles(self):
        with open(self.feedbackDatabase, 'r') as feedbackDatabaseFile:
            
            databaseToRead = json.load(feedbackDatabaseFile)
            feedbackDatabaseFile.close()


        for basicSystem in databaseToRead['feedbacks']['systems']:
            
            for specificSystem in databaseToRead['feedbacks']['systems'][basicSystem]:

                    feedbackFileToWrite = Document()

                    self.createFeedbackDir()
                    self.createSystemDir()

                    try:
                        relBasicSystemPath = f'../../Feedbacks/Systems/{basicSystem}'
                        os.mkdir(self.relativeFilepathToAbsolute(relBasicSystemPath))
                        print('Dir created', relBasicSystemPath)
                    except FileExistsError:
                        print(f'Directory {relBasicSystemPath} already exists')
                    finally:
                        feedbackFilePathRel = f'../../Feedbacks/Systems/{basicSystem}/{specificSystem}.docx'
                        feedbackFilePathAbs = self.relativeFilepathToAbsolute(feedbackFilePathRel)    


                    

                    feedbackFileToWrite.add_heading(f'{specificSystem}', 0)
                    
                    for ship in databaseToRead['feedbacks']['systems'][basicSystem][specificSystem]:
                        
                        feedbackFileToWrite.add_heading(f'{ship}:', level=1)
                        
                        questionAnswerTable = feedbackFileToWrite.add_table(rows=1, cols=2)
                        questionAnswerTable.style = 'Table Grid'
                        questionAnswerTable.autofit = False
                        
                        headingCells = questionAnswerTable.rows[0].cells
                        headingCells[0].text = 'Question'
                        headingCells[1].text = 'Answers'
                        headingCells[0].width = Inches(1.5)
                        headingCells[1].width = Inches(5.0)    
                    
                        for question in databaseToRead['feedbacks']['systems'][basicSystem][specificSystem][ship]:
                            
                            if question in self.questionToExclude:
                                continue

                            row_cells = questionAnswerTable.add_row().cells

                            question = self.replaceStrings(question, {'\t': ' '})
                            row_cells[0].text = question
                            row_cells[0].width = Inches(1.5)

                            for answer in databaseToRead['feedbacks']['systems'][basicSystem][specificSystem][ship][question]:
                                
                                answer = self.replaceStrings(answer, {'\t': ' '})

                                ansToAdd = f'{self.mode}: {answer}'

                                row_cells[1].add_paragraph(ansToAdd, style='List Bullet')
                                row_cells[1].width = Inches(5.0)

                    feedbackFileToWrite.save(feedbackFilePathAbs)
        

    def main(self):
        self.createCsvDir()
        self.csvSystemsFeedbackFile = self.convertXlsxToCsv(self.excelSystemsFeedackFile)
        self.listSystemsFeedbackFile = self.convertCsvToList(self.csvSystemsFeedbackFile)
        self.updateSystemsDatabase()
        self.createFeedbackFiles()

'''if __name__ == '__main__':
    excelFilepath = "NB518_design.xlsx"
    
    systemsUpdate = SystemsUpdate(excelFilepath, '../databases/feedbackDatabaseTest.json', 'Design')
    systemsUpdate.main()
'''