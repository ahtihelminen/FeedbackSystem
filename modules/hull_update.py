from tools import Tools
import json
from docx import Document
from docx.shared import Inches
import os

class HullUpdate(Tools):
    def __init__(self, excelHullFeedackFile, feedbackDatabase, ship):
        super().__init__()
        self.excelHullFeedackFile = excelHullFeedackFile
        self.feedbackDatabase = feedbackDatabase
        self.csvHullFeedbackFile = None
        self.listHullFeedbackFile = None
        self.ship = ship
        self.questionToExclude = ['ID', 'Start time', 'Completion time', 'Email', 'Name', 'valitse littera', 'choose system code', 'choose your area', '1000 Ship general design', '3000 Hull', '4000 Interior', '5000 HVAC', '6000 Propulsion', '7000 Machinery', '8000 Deck', '9000 Electric']
        self.mode = None

    def applicableString(self, stringToChange):
        stringToChange = stringToChange.strip()
        stringToChange = stringToChange.replace('\t', ' ')
        stringToChange = stringToChange.replace('\n', ' ')
        return stringToChange


    def updateHullDatabase(self):

        self.initializeDatabase()

        self.feedbackDatabase = self.relativeFilepathToAbsolute(self.feedbackDatabase)
        questions = self.listHullFeedbackFile[0]
        database = self.readJSON(self.feedbackDatabase)


        with open(self.feedbackDatabase, 'w') as databaseToWrite:
            try:
                
                databaseToUpdate = database
                
                for row in self.listHullFeedbackFile[1:]:
                    
                    questionAnswerDict = self.questionAnswerDict(row, questions)
                    questionAnswerDict = self.removeQuestionsWithNoAnswer(questionAnswerDict)
                    area = questionAnswerDict['choose your area']
                    
                    if area not in databaseToUpdate['feedbacks']['hull']:
                        databaseToUpdate['feedbacks']['hull'][area] = {}

                    if self.ship not in databaseToUpdate['feedbacks']['hull'][area]:
                        databaseToUpdate['feedbacks']['hull'][area][self.ship] = {}

                    for question, answer in questionAnswerDict.items():

                        if question not in databaseToUpdate['feedbacks']['hull'][area][self.ship]:
                            databaseToUpdate['feedbacks']['hull'][area][self.ship][question] = []
                        
                        if answer not in databaseToUpdate['feedbacks']['hull'][area][self.ship][question]:
                            databaseToUpdate['feedbacks']['hull'][area][self.ship][question].append(answer)

                json.dump(databaseToUpdate, databaseToWrite, indent=4)
            except Exception as e:
                print('Error in updateDatabase()', e)
                json.dump(database, databaseToWrite, indent=4)
                raise TypeError('Incorrect file or mode selected!')
        databaseToWrite.close()


    def createFeedbackFiles(self):
        with open(self.feedbackDatabase, 'r') as feedbackDatabaseFile:
            
            databaseToRead = json.load(feedbackDatabaseFile)
            feedbackDatabaseFile.close()


        for hullPart in databaseToRead['feedbacks']['hull']:

            feedbackFileToWrite = Document()

            self.createFeedbackDir()
            self.createHullDir()

            try:
                relBasicSystemPath = f'../Feedbacks/Hull/{hullPart}'
                os.mkdir(self.relativeFilepathToAbsolute(relBasicSystemPath))
                print('Dir created', relBasicSystemPath)
            except FileExistsError:
                print(f'Directory {relBasicSystemPath} already exists')
            finally:
                feedbackFilePathRel = f'../feedbacks/hull/{hullPart}.docx'
                feedbackFilePathAbs = self.relativeFilepathToAbsolute(feedbackFilePathRel)

            feedbackFileToWrite.add_heading(f'{hullPart}', 0)
            
            for ship in databaseToRead['feedbacks']['hull'][hullPart]:
                
                feedbackFileToWrite.add_heading(f'{ship}:', level=1)

                questionAnswerTable = feedbackFileToWrite.add_table(rows=1, cols=2)
                questionAnswerTable.style = 'Table Grid'
                questionAnswerTable.autofit = False
                
                headingCells = questionAnswerTable.rows[0].cells
                headingCells[0].text = 'Question'
                headingCells[1].text = 'Answers'
                headingCells[0].width = Inches(1.5)
                headingCells[1].width = Inches(5.0)    
            
                for question in databaseToRead['feedbacks']['hull'][hullPart][ship]:
                    
                    if question in self.questionToExclude:
                        continue
                    
                    if question == 'Valitse / choose':
                        self.mode = databaseToRead['feedbacks']['hull'][hullPart][ship][question][0]
                        continue

                    
                    row_cells = questionAnswerTable.add_row().cells

                    question = self.replaceStrings(question, {'\t': ' '})
                    row_cells[0].text = question
                    row_cells[0].width = Inches(1.5)

                    for answer in databaseToRead['feedbacks']['hull'][hullPart][ship][question]:
                        answer = self.replaceStrings(answer, {'\t': ' '})
                        


                        ansToAdd = f'{self.mode}: {answer}'

                        row_cells[1].add_paragraph(ansToAdd, style='List Bullet')
                        row_cells[1].width = Inches(5.0)

            feedbackFileToWrite.save(feedbackFilePathAbs)


    def test(self):
        questions = self.listHullFeedbackFile[0]
        qadict = self.questionAnswerDict(self.listHullFeedbackFile[1], questions)
        print(self.removeQuestionsWithNoAnswer(qadict))
        print(self.extracthull(qadict))
        dbPath = self.relativeFilepathToAbsolute('../databases/feedbackDatabaseTest.json')
        database = self.readJSON(dbPath)
        print(database['feedbacks']['hull'])


    def main(self):
        self.createCsvDir()
        self.csvHullFeedbackFile = self.convertXlsxToCsv(self.excelHullFeedackFile)
        self.listHullFeedbackFile = self.convertCsvToList(self.csvHullFeedbackFile)
        #self.test()
        self.updateHullDatabase()
        self.createFeedbackFiles()


'''if __name__ == '__main__':
    excelFilepath = "NB518_hull.xlsx"
    
    hullUpdate = HullUpdate(excelFilepath, '../databases/feedbackDatabaseTest.json')
    hullUpdate.main()
'''