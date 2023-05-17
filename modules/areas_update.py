from tools import Tools
import json
from docx import Document
from docx.shared import Inches

class AreasUpdate(Tools):
    def __init__(self, excelAreasFeedackFile, feedbackDatabase, ship):
        super().__init__()
        self.excelAreasFeedackFile = excelAreasFeedackFile
        self.feedbackDatabase = feedbackDatabase
        self.csvAreasFeedbackFile = None
        self.listAreasFeedbackFile = None
        self.ship = ship
        self.questionToExclude = [
            'ID',
            'Start time',
            'Completion time',
            'Email',
            'Name',
            'Valitse alue / alueet',
            'valitse littera',
            'choose system code',
            '1000 Ship general design',
            '3000 Hull',
            '4000 Interior',
            '5000 HVAC',
            '6000 Propulsion',
            '7000 Machinery',
            '8000 Deck',
            '9000 Electric',
            'Valitse. Oletko?',
            'Oliko alueella sinun vastuullasi alihankintaa'
        ]
        self.mode = None


    def applicableString(self, stringToChange):
        
        stringToChange = stringToChange.replace('\t', ' ')
        stringToChange = stringToChange.replace('\n', ' ')
        stringToChange = stringToChange.strip()

        print(stringToChange)
        return stringToChange


    def extractAreasForCreate(self, Q_A_dict):
        try:
            if 'Valitse alue / alueet' not in Q_A_dict:
                return ['']
            areasList = Q_A_dict['Valitse alue / alueet'].split(';')
            areaDirList = []

            dirDict = {
                'H': 'Hissit',
                'K': 'Kuilut',
                'P': 'Portaikot'
            }

            for area in areasList:
                
                singleAreaSplit = self.applicableString(area).split(' ')
                
                if len(singleAreaSplit) == 2:
                
                    if singleAreaSplit[0][0] in dirDict:
                        areaDirList.append(dirDict[singleAreaSplit[0][0]])
                
                    else:
                        areaDirList.append(singleAreaSplit[1])

                elif len(singleAreaSplit) == 3:
                    
                    if singleAreaSplit[2] in ['1', '2', '3', '4', '1,', '2,', '3,', '4,']:
                        areaDirList.append(singleAreaSplit[1])
                    
                    else:
                        areaDirList.append(f'{singleAreaSplit[1]} {singleAreaSplit[2]}')
                
                else:
                    areaDir = ''
                    for i in range(1, len(singleAreaSplit)-1):
                        areaDir += f'{singleAreaSplit[i]} '
                    areaDirList.append(areaDir)
            
            return areaDirList
        except Exception as e:
            print('Error in extractAreasForCreate()', e)
            raise TypeError('Incorrect file or mode selected!')


    def updateAreasDatabase(self):
        

        self.initializeDatabase()

        self.feedbackDatabase = self.relativeFilepathToAbsolute(self.feedbackDatabase)
        questions = self.listAreasFeedbackFile[0]
        database = self.readJSON(self.feedbackDatabase)

        

        with open(self.feedbackDatabase, 'w') as databaseToWrite:
            try:
                
                databaseToUpdate = database
                
                for row in self.listAreasFeedbackFile[1:]:
                    
                    questionAnswerDict = self.questionAnswerDict(row, questions)
                    questionAnswerDict = self.removeQuestionsWithNoAnswer(questionAnswerDict)
                    self.QADict = questionAnswerDict

                    areasList = self.extractAreasForCreate(questionAnswerDict)


                    for specificArea in areasList:
                        if specificArea == '':
                            continue

                        specificArea = self.applicableString(specificArea)
                        
                        if specificArea not in databaseToUpdate['feedbacks']['areas']:
                            databaseToUpdate['feedbacks']['areas'][specificArea] = {}

                        if self.ship not in databaseToUpdate['feedbacks']['areas'][specificArea]:
                            databaseToUpdate['feedbacks']['areas'][specificArea][self.ship] = {}

                        for question, answer in questionAnswerDict.items():

                            if question not in databaseToUpdate['feedbacks']['areas'][specificArea][self.ship]:
                                databaseToUpdate['feedbacks']['areas'][specificArea][self.ship][question] = []
                            
                            if answer not in databaseToUpdate['feedbacks']['areas'][specificArea][self.ship][question]:
                               databaseToUpdate['feedbacks']['areas'][specificArea][self.ship][question].append(answer)

                json.dump(databaseToUpdate, databaseToWrite, indent=4)
            except Exception as e:
                print('Error in updateDatabase()', e)
                json.dump(database, databaseToWrite, indent=4)
                raise TypeError('Incorrect file or mode selected!')
        databaseToWrite.close()


    def createFeedbackFiles(self):
        
        self.createFeedbackDir()
        self.createAreaDir()

        with open(self.feedbackDatabase, 'r') as feedbackDatabaseFile:
            
            databaseToRead = json.load(feedbackDatabaseFile)
            feedbackDatabaseFile.close()


        for area in databaseToRead['feedbacks']['areas']:
            
            feedbackFileToWrite = Document()

            
            feedbackFilePathRel = f'../../Feedbacks/Areas/{area}.docx'
            feedbackFilePathAbs = self.relativeFilepathToAbsolute(feedbackFilePathRel)

            feedbackFileToWrite.add_heading(f'{area}', 0)

            for ship in databaseToRead['feedbacks']['areas'][area]:
                
                feedbackFileToWrite.add_heading(f'{ship}:', level=1)
                
                questionAnswerTable = feedbackFileToWrite.add_table(rows=1, cols=2)
                questionAnswerTable.style = 'Table Grid'
                questionAnswerTable.autofit = False
                
                headingCells = questionAnswerTable.rows[0].cells
                headingCells[0].text = 'Question'
                headingCells[1].text = 'Answers'
                headingCells[0].width = Inches(1.5)
                headingCells[1].width = Inches(5.0)    
            
                for question in databaseToRead['feedbacks']['areas'][area][ship]:
                    
                    if question == 'Valitse. Oletko?':
                        self.mode = databaseToRead['feedbacks']['areas'][area][ship][question][0]

                    if question in self.questionToExclude:
                        continue

                    

                    row_cells = questionAnswerTable.add_row().cells

                    question = self.replaceStrings(question, {'\t': ' '})
                    row_cells[0].text = question
                    row_cells[0].width = Inches(1.5)

                    for answer in databaseToRead['feedbacks']['areas'][area][ship][question]:
                        
                        answer = self.replaceStrings(answer, {'\t': ' '})

                        ansToAdd = f'{self.mode}: {answer}'

                        row_cells[1].add_paragraph(ansToAdd, style='List Bullet')
                        row_cells[1].width = Inches(5.0)
            
            feedbackFileToWrite.save(feedbackFilePathAbs)


    def test(self):
        questions = self.listAreasFeedbackFile[0]
        qadict = self.questionAnswerDict(self.listAreasFeedbackFile[1], questions)
        print(self.removeQuestionsWithNoAnswer(qadict))
        dbPath = self.relativeFilepathToAbsolute('../databases/feedbackDatabaseTest.json')
        database = self.readJSON(dbPath)
        print(database['feedbacks']['Areas'])


    def main(self):
        self.csvAreasFeedbackFile = self.convertXlsxToCsv(self.excelAreasFeedackFile)
        self.listAreasFeedbackFile = self.convertCsvToList(self.csvAreasFeedbackFile)
        #self.test()
        self.updateAreasDatabase()
        self.createFeedbackFiles()

if __name__ == '__main__':
    excelFilepath = "../feedbacksExcel/NB518_areas.xlsx"
    
    AreasUpdate = AreasUpdate(excelFilepath, '../databases/feedbackDatabaseTest.json')
    AreasUpdate.main()
